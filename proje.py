import os
import sys
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from browser_use import Agent
import asyncio
from docx import Document

load_dotenv()

print("Python version:", sys.version)
print("Current working directory:", os.getcwd())
print("GOOGLE_APPLICATION_CREDENTIALS:", os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"))
print("GEMINI_API_KEY status:", "Set" if os.environ.get("GEMINI_API_KEY") else "Not Set")

def read_list_from_docx(filepath):
    try:
        doc = Document(filepath)
        return [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return []

disorders = read_list_from_docx('disorders.docx')

TOPIC_QUERIES = {
    "therapy": "therapy methods used to treat {disorder}",
    "medication": "medication used for {disorder}",
    "symptoms": "symptoms of {disorder}",
    "diagnosis": "how to diagnose {disorder}",
    "related_disorders": "disorders related to {disorder}",
    "diet": "diet impact on {disorder}",
    "lifestyle_positive": "lifestyles with positive impact on {disorder}",
    "lifestyle_negative": "lifestyles with negative impact on {disorder}",
    "sports": "sports activities beneficial for {symptom}"
}

def generate_query(topic, disorder, symptom=None):
    if topic == "sports" and symptom:
        return TOPIC_QUERIES.get(topic, '').format(symptom=symptom)
    if '{disorder}' in TOPIC_QUERIES.get(topic, ''):
         return TOPIC_QUERIES.get(topic, '').format(disorder=disorder)
    return TOPIC_QUERIES.get(topic, '')

def extract_symptoms_from_summary(summary):
    print(f"--- Symptoms Summary (for Extraction) ---")
    print(summary)
    print(f"-----------------------------------------")
    print("NOTE: The symptoms shown above might be an error message if the search failed.")
    print("      You need to implement extract_symptoms_from_summary to parse actual symptom text.")

    symptoms_list = []
   
    if isinstance(summary, str) and "Error processing symptoms" not in summary:
       
        possible_symptom_keywords = ["anxiety", "depression", "fatigue", "pain", "insomnia", "irritability", "low mood"] 
        for keyword in possible_symptom_keywords:
            if keyword in summary.lower():
                
                 if keyword.capitalize() not in symptoms_list: 
                     symptoms_list.append(keyword.capitalize())

    if not symptoms_list and isinstance(summary, str) and "Error processing symptoms" not in summary and len(summary) > 50: 
         print("Warning: Placeholder symptom extraction found no keywords. Consider improving implementation.")
        


    if not symptoms_list:
        print("Warning: No symptoms extracted. Skipping sports query for this disorder.")

    return symptoms_list


def summarize_with_llm(llm, result_content, topic, disorder, symptom=None):
    if not isinstance(result_content, str):
        print(f"Warning: result_content for {topic} - {disorder} is not a string ({type(result_content)}). Converting to string.")
        result_content = str(result_content)

    if "Error processing" in result_content or ">=" in result_content or "unexpected keyword argument" in result_content:
        print(f"Skipping LLM summarization for {topic} due to search error: {result_content}")
        return result_content 

    if symptom:
        prompt = f"""
        ---Articles---
        {result_content}

        Reduce detailed summaries about activities beneficial for the symptom '{symptom}' related to {disorder}, based on the articles' context.
        Write a comprehensive, well-structured report for this symptom, at least 200 words (adjust word count as needed). Focus specifically on the symptom.
        """
    else:
        prompt = f"""
        ---Articles---
        {result_content}

        Summarize information about {topic.replace('_', ' ')} for {disorder}.
        Write a comprehensive, well-structured report, at least 300 words (adjust word count as needed).
        """
    try:
        print(f"Calling LLM to summarize {topic} for {disorder} (Symptom: {symptom})...")
        summary = llm.invoke(prompt)
        print("LLM call successful.")
        if hasattr(summary, 'content'):
             return summary.content
        return str(summary)
    except Exception as e:
        print(f"Error calling LLM for summary of {topic} - {disorder}: {e}")
        return f"Error generating summary for {topic} - {disorder}: {e}"


def disorder_template():
    return {
        "therapy": [],
        "medication": "",
        "symptoms": "",
        "diagnosis": "",
        "related_disorders": "",
        "diet": "",
        "lifestyle_positive": "",
        "lifestyle_negative": "",
        "sports": {}
    }

reports = {}

def write_disorder_doc(disorder, data, filename):
    doc = Document()
    doc.add_heading(disorder, 0)

    for topic, content in data.items():
        if topic == "sports":
            doc.add_heading("Sports Activities", level=1)
            if isinstance(content, dict):
                if not content:
                    doc.add_paragraph("No specific sports activities found for symptoms.")
                for symptom, sports_content in content.items():
                    doc.add_heading(f"Beneficial for {symptom}", level=2)
                    doc.add_paragraph(sports_content)
            else:
                 doc.add_paragraph(str(content))
        elif topic == "therapy":
             doc.add_heading("Therapy Methods", level=1)
             if isinstance(content, list):
                 if not content:
                     doc.add_paragraph("No therapy methods found.")
                 for i, therapy_summary in enumerate(content):
                     doc.add_heading(f"Therapy Summary {i+1}", level=2)
                     doc.add_paragraph(therapy_summary)
             else:
                 doc.add_paragraph(str(content))
        else:
            doc.add_heading(topic.replace('_', ' ').capitalize(), level=1)
            doc.add_paragraph(content)

    try:
        os.makedirs("outputs", exist_ok=True)
        safe_filename = "".join([c for c in filename if c.isalnum() or c in (' ', '-', '_', '(' , ')')]).rstrip() 
            safe_filename = "report"
        filepath = os.path.join("outputs", f"{safe_filename}.docx")
        doc.save(filepath)
        print(f"Report saved to {filepath}")
    except Exception as e:
        print(f"Error writing Word document for {disorder}: {e}")

async def main():
    gemini_api_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_api_key:
        print("Error: GEMINI_API_KEY not found in environment variables or .env file.")
        print("Please set the GEMINI_API_KEY variable.")
        return

    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash-latest", api_key=gemini_api_key)

    agent = Agent(
        task="search in internet",
        llm=llm,
    )
    

    disorders = read_list_from_docx('disorders.docx')
    if not disorders:
        print("Error: No disorders found in disorders.docx. Please add disorders to the file.")
        return

    for disorder in disorders:
        print(f"\n--- Processing Disorder: {disorder} ---")
        reports[disorder] = disorder_template()

        for topic in TOPIC_QUERIES.keys():
            if topic == "sports":
                continue

            query = generate_query(topic, disorder)
            if not query:
                 print(f"Skipping {topic} for {disorder} due to empty query.")
                 reports[disorder][topic] = "Query could not be generated."
                 continue

            print(f"Searching for: {query}")
            try:
                print("--- Using DUMMY search result to test summarization and document writing ---")
                result = f"DUMMY SEARCH RESULT for '{query}'. Replace this with real search results.\n\nExample research summary 1 about {topic} for {disorder}: This study found positive outcomes...\nExample research summary 2: Another paper discussed...\nSymptom keywords for {disorder}: fatigue, pain, insomnia."
                print(f"Dummy result: {result[:200]}...")
                
                summary = summarize_with_llm(llm, result, topic, disorder)

                if topic == "therapy":
                    reports[disorder]["therapy"].append(summary)
                else:
                    reports[disorder][topic] = summary
            except Exception as e:
                print(f"Error during processing {topic} for {disorder}: {type(e).__name__}: {e}")
                reports[disorder][topic] = f"Error processing {topic}: {e}"

        print(f"\nExtracting symptoms for {disorder}...")
        symptoms_summary_content = reports[disorder]["symptoms"]
        if not isinstance(symptoms_summary_content, str):
             symptoms_summary_content = str(symptoms_summary_content)

        symptoms = extract_symptoms_from_summary(symptoms_summary_content) 

        if symptoms:
            print(f"Extracted symptoms: {symptoms}")
            for symptom in symptoms:
                print(f"Processing sports for symptom: {symptom}")
                sports_query = generate_query("sports", disorder, symptom=symptom)
                if not sports_query:
                     print(f"Skipping sports query for symptom '{symptom}' due to empty query.")
                     reports[disorder]["sports"][symptom] = "Query could not be generated."
                     continue

                print(f"Searching for: {sports_query}")
                try:
                    print("--- Using DUMMY search result for SPORTS to test summarization and document writing ---")
                    sports_result = f"DUMMY SEARCH RESULT for '{sports_query}'. Replace this with real search results.\n\nResearch indicates exercise is beneficial for {symptom}.\nOne paper suggests activities like swimming...\nAnother study mentions walking helps with {symptom}."
                    print(f"Dummy result: {sports_result[:200]}...")
                
                    sports_summary = summarize_with_llm(llm, sports_result, "sports", disorder, symptom=symptom)
                    reports[disorder]["sports"][symptom] = sports_summary

                except Exception as e:
                     print(f"Error during processing sports for {symptom}: {type(e).__name__}: {e}")
                     reports[disorder]["sports"][symptom] = f"Error processing sports for {symptom}: {e}"
        else:
             print(f"No symptoms extracted for {disorder} or symptom extraction failed. Skipping sports section.")


        write_disorder_doc(disorder, reports[disorder], f"{disorder}_report")


if __name__ == "__main__":
    try:
        loop = asyncio.get_running_loop()
        if loop.is_running():
             print("Running in existing asyncio loop.")
             loop.create_task(main())
        else:
             asyncio.run(main())
    except RuntimeError:
        asyncio.run(main())
